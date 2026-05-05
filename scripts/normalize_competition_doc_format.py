from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "deliverables" / "competition_doc_2026" / "智能拍照助手-物联网应用类作品技术文档-2026版.docx"
ALT_DOC_PATH = ROOT / "deliverables" / "competition_doc_2026" / "智能拍照助手-物联网应用类作品技术文档-2026版-格式修正版.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)
SONG = "宋体"
HEADING_RE = re.compile(r"^\d+(?:\.\d+)*\.?\s")
BULLET_RE = re.compile(r"^\s*[•●▪◦·-]\s*")


def set_run_font(run, *, size_pt: float, bold: bool) -> None:
    run.font.name = SONG
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), SONG)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), SONG)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), SONG)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = BLACK


def set_cell_shading_white(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is not None:
        tc_pr.remove(shd)
    tc_pr.append(parse_xml(rf'<w:shd {nsdecls("w")} w:fill="FFFFFF"/>'))


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders_black(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)
    tbl_pr.append(borders)


def normalize_paragraph(paragraph) -> None:
    text = paragraph.text.strip()
    if not text:
        return

    # Remove list formatting and leading bullet glyphs so text can be copied directly.
    paragraph.style = None
    if BULLET_RE.match(text):
        clean = BULLET_RE.sub("", text)
        paragraph.clear()
        paragraph.add_run(clean)
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = None

    stripped = paragraph.text.strip()
    is_heading = bool(HEADING_RE.match(stripped))
    size = 14 if is_heading else 10.5
    bold = is_heading

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if not is_heading else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run, size_pt=size, bold=bold)


def normalize_tables(doc: Document) -> None:
    for table in doc.tables:
        set_table_borders_black(table)
        for r, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell)
                set_cell_shading_white(cell)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        set_run_font(run, size_pt=10.5, bold=(r == 0))


def normalize_headers_footers(doc: Document) -> None:
    for section in doc.sections:
        for part in (section.header, section.footer):
            for para in part.paragraphs:
                if not para.text.strip():
                    continue
                for run in para.runs:
                    set_run_font(run, size_pt=9, bold=False)


def main() -> None:
    doc = Document(str(DOC_PATH))

    for para in doc.paragraphs:
        normalize_paragraph(para)

    normalize_tables(doc)
    normalize_headers_footers(doc)

    try:
        doc.save(str(DOC_PATH))
        print(DOC_PATH)
    except PermissionError:
        doc.save(str(ALT_DOC_PATH))
        print(ALT_DOC_PATH)


if __name__ == "__main__":
    main()
