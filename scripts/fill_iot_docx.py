from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"E:\BianChengWenJian\Camera Assistant")
TEMPLATE = ROOT / "docs" / "技术文档模板.docx"
OUTPUT = ROOT / "docs" / "云影随行-物联网应用类作品技术文档-正式版.docx"


def set_run_font(run, name: str, size_pt: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table):
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.15
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_run_font(run, "宋体", 10.5)
            if row_idx == 0:
                set_cell_shading(cell, "EAF4F8")
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        set_run_font(run, "微软雅黑", 10.5, bold=True)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    style: str = "Normal",
    bold: bool = False,
    center: bool = False,
    first_line_chars: float | None = 2.0,
    size: float = 12,
    font_name: str = "宋体",
    color: str | None = None,
    space_after: float = 6,
) -> None:
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, font_name, size, bold=bold, color=color)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if first_line_chars is None:
            p.paragraph_format.first_line_indent = Pt(0)
        else:
            p.paragraph_format.first_line_indent = Cm(0.74 * first_line_chars)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        bullet = p.add_run("• ")
        set_run_font(bullet, "微软雅黑", 12, bold=True, color="C98A2E")
        text = p.add_run(item)
        set_run_font(text, "宋体", 12)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, "微软雅黑", 12, bold=True)


def clear_after_index(doc: Document, keep_count: int) -> None:
    paras = list(doc.paragraphs)
    for p in paras[keep_count:]:
        elem = p._element
        elem.getparent().remove(elem)


def update_cover(doc: Document) -> None:
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("作品编号"):
            p.text = "作品编号：待赛事系统分配"
        elif txt.startswith("作品名称"):
            p.text = "作品名称：云影随行智能拍照助手"
        elif txt.startswith("版本编号"):
            p.text = "版本编号：V1.0 正式版"
        elif txt.startswith("填写日期"):
            p.text = "填写日期：2026年4月28日"
        for run in p.runs:
            set_run_font(run, "宋体", 12)


def apply_styles(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18)
        sec.right_margin = Cm(3.18)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(14)
    h1.font.bold = True


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, "黑体", 14, bold=True)


def add_rows(table, rows: list[list[str]]):
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def build_doc():
    doc = Document(str(TEMPLATE))
    apply_styles(doc)
    update_cover(doc)
    clear_after_index(doc, 9)

    add_heading(doc, "作品概述")
    add_paragraph(
        doc,
        "云影随行智能拍照助手是一套面向数字生活拍照场景的多端协同系统，主要解决自拍、多人合照、旅游打卡和模板复拍等场景中不会构图、抓不住时机、依赖他人帮拍、拍后难以复盘等实际问题。"
        "系统不是单一拍照 App，而是由移动端、边缘设备运行时、云端业务后端和管理后台共同组成。移动端负责用户交互、本机拍摄与设备联动；边缘端负责实时视觉处理、模板构图、云台控制与抓拍执行；"
        "云端负责账户、模板、历史和 AI 任务管理；后台负责运维配置和业务管理。"
    )
    add_paragraph(
        doc,
        "本作品的应用领域包括智能拍照辅助、物联网终端联动、人机交互、视觉感知与边缘智能。与仅提供滤镜、美颜或简单快门控制的产品不同，本作品强调拍前引导、拍中执行、拍后分析、历史沉淀的完整闭环。"
        "当前源码已经实现用户登录注册、移动端独立拍摄、设备联动推流、设备预览回传、云台控制、手势抓拍、AI 图片分析、模板管理和后台管理等核心能力。"
    )
    add_paragraph(doc, "本作品的创新性主要体现在以下四个方面：", first_line_chars=None)
    add_bullets(doc, [
        "设计思路创新：将拍照过程拆解为拍前引导、拍中执行、拍后复盘三个阶段，通过多端协同形成可闭环的智能拍照流程。",
        "技术实现创新：在同一项目中同时实现移动端、边缘视觉、云端业务和后台管理四个运行平面，并通过 REST、WebSocket 和保留的 WebRTC 能力组织协作。",
        "硬件联动创新：将人体检测、姿态辅助、手势触发与云台控制结合，使视觉理解结果能够进一步驱动真实设备动作。",
        "应用场景创新：针对自拍、多人合照和模板复拍等高频场景，提供比传统相机或单一云台 App 更完整的辅助拍摄能力。"
    ])

    add_heading(doc, "需求分析")
    add_paragraph(
        doc,
        "当前数字生活场景中，拍照已经成为社交分享、旅游留念、家庭纪念和个人内容创作的基础能力。但从用户真实体验看，拍照仍然存在明显痛点："
        "一是普通用户缺少稳定的构图能力，人物位置、角度和留白常常处理不好；二是自拍和多人合照场景下抓拍时机难把握，经常需要找路人帮拍或反复试拍；"
        "三是拍完之后缺少统一的历史沉淀和复拍经验复用；四是现有产品大多只覆盖某一个局部环节，例如只负责拍照、只负责云台控制或只负责 AI 修图，难以形成闭环体验。"
    )
    add_paragraph(
        doc,
        "本作品面向的核心用户包括有自拍、旅游打卡、家庭合照需求的普通手机用户，希望减少依赖他人帮拍的用户，以及对模板复拍、智能构图、设备辅助抓拍有需求的轻度内容创作者。"
        "从当前源码所面向的问题看，本作品并非简单对标某一单一竞品，而是综合吸收原生相机、美颜拍照应用、云台配套 App 和 AI 修图工具的部分能力，再以多端协同方式形成差异化。"
    )
    add_subtitle(doc, "竞品分析表")
    table = doc.add_table(rows=1, cols=6)
    for i, h in enumerate(["对比维度", "原生相机", "美颜/拍照类 App", "云台配套 App", "AI 修图工具", "本作品"]):
        table.cell(0, i).text = h
    add_rows(table, [
        ["本机拍照", "强", "强", "中", "弱", "强"],
        ["实时构图引导", "弱", "中", "中", "弱", "强"],
        ["设备云台联动", "无", "无", "强", "无", "强"],
        ["手势触发抓拍", "弱", "弱", "中", "无", "强"],
        ["拍后 AI 分析", "弱", "中", "弱", "强", "强"],
        ["历史沉淀与复拍", "中", "中", "弱", "中", "强"],
        ["多端协同闭环", "弱", "弱", "中", "弱", "强"],
    ])
    style_table(table)
    add_paragraph(
        doc,
        "从竞品对比可以看出，本作品的核心差异不在单个功能点，而在于移动端交互、边缘设备执行、云端业务管理与 AI 分析组合形成的整体能力。"
        "这一定义与当前仓库的实际代码结构完全一致。"
    )

    add_heading(doc, "技术方案")
    add_subtitle(doc, "1. 总体技术路线")
    add_paragraph(
        doc,
        "系统采用移动端、边缘设备、云端后端和管理后台分层协同的技术路线。移动端模块基于 Flutter 实现，负责登录注册、本机拍摄、模板管理、历史查看和设备联动；"
        "边缘运行时模块基于 FastAPI、OpenCV、MediaPipe 构建，负责接收手机推流、进行人体与手势检测、生成预览反馈、控制云台并完成本地抓拍；"
        "云端后端模块基于 FastAPI 与 SQLAlchemy 构建，负责用户、套餐、模板、抓拍历史、AI 任务和 AI Provider 配置管理；"
        "管理后台模块基于 Vue 3、Element Plus、Pinia 构建，负责用户、设备、推荐模板、抓拍记录和 AI 配置维护。数据层使用 PostgreSQL 保存结构化数据，使用后端上传目录和设备本地抓拍目录分别保存后端上传图片与设备本地抓拍结果。"
    )
    add_paragraph(
        doc,
        "从通信方式看，移动端与业务后端主要通过 REST 交互，移动端与边缘设备之间主要通过 REST 与 WebSocket 交互，并保留 WebRTC 联动能力。"
        "这样的划分使系统能够在局域网设备联动场景中保持较低时延，同时将业务管理和数据沉淀集中到云端处理。"
    )

    add_subtitle(doc, "2. 硬件组成与来源")
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["硬件组成", "作用", "源码对应", "来源说明"]):
        table.cell(0, i).text = h
    add_rows(table, [
        ["Android 智能手机", "运行移动端 App，采集相机图像并作为用户交互入口", "移动端依赖声明与相机采集代码", "通用商购设备，演示机型可在答辩时补充"],
        ["树莓派 4B/5 或同类 Linux 边缘主机", "运行设备运行时，完成实时视觉处理和设备控制", "设备运行时配置与树莓派性能档", "当前代码按树莓派运行档位设计"],
        ["双轴 TTL 总线舵机云台", "执行水平与俯仰角度控制", "云台控制器与串口总线驱动实现", "通用商购模块，可替换为同类总线舵机方案"],
        ["USB 转串口/TTL 控制链路", "连接边缘设备与舵机控制器", "串口驱动默认参数为 /dev/ttyUSB0、115200 波特率", "通用商购模块"],
        ["电源与局域网/热点环境", "提供供电与局域网通信能力", "代码未强绑定具体品牌型号", "实物部署时按现场条件配置"],
    ])
    style_table(table)
    add_paragraph(
        doc,
        "当前代码同时支持模拟驱动与 TTL 总线驱动两种云台控制模式，说明系统在开发阶段兼顾了无实物调试和真实硬件联调两种场景。"
        "正式提交时如有条件，建议附上实物接线图、供电说明和安装照片。"
    )

    add_subtitle(doc, "3. 软件组成与关键依赖")
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["软件模块", "主要技术栈", "说明"]):
        table.cell(0, i).text = h
    add_rows(table, [
        ["移动端", "Flutter、camera、http、shared_preferences、姿态检测、WebRTC", "负责用户交互、本机拍摄、设备联动和历史查看"],
        ["边缘端", "FastAPI、OpenCV、MediaPipe、NumPy、Pillow、pyserial", "负责推流处理、实时检测、模板构图、云台控制和本地抓拍"],
        ["云端后端", "FastAPI、SQLAlchemy、psycopg、httpx", "负责鉴权、业务模型、历史记录、AI 任务和 Provider 编排"],
        ["管理后台", "Vue 3、Vite、Element Plus、Pinia、Axios", "负责用户、设备、模板、Provider 等管理"],
        ["数据库", "PostgreSQL", "保存用户、设备、模板、会话、抓拍和 AI 任务等结构化数据"],
    ])
    style_table(table)

    add_subtitle(doc, "4. 接口通用性、扩展性与代码规范")
    add_bullets(doc, [
        "后端接口按移动端与管理端分组，移动端接口与管理端接口边界清晰。",
        "设备运行时接口按会话、推流、状态、控制、抓拍、模板、AI 和 WebRTC 等能力拆分，结构清晰。",
        "移动端通过独立服务层封装业务后端调用、设备运行时调用和联动通信能力，便于维护和扩展。",
        "云台控制采用抽象驱动接口，下层既支持模拟驱动，也支持 TTL 总线实物驱动，硬件替换成本较低。",
        "AI 能力通过后端编排与设备侧辅助逻辑分层组织，便于后续切换模型或增加 Provider。"
    ])

    add_subtitle(doc, "5. 硬件设计合理性与安全性")
    add_paragraph(
        doc,
        "从当前代码看，硬件控制层已经考虑了角度边界约束、平滑移动、反馈刷新和超时控制等机制。"
        "例如云台角度通过最小角、最大角和回中角进行限制，避免指令越界；平滑移动采用分步逼近而不是一次性跳变，降低机械抖动风险；串口驱动提供超时设置与反馈刷新机制，便于状态监测。"
    )
    add_paragraph(
        doc,
        "由于仓库中尚未包含完整的接线图和供电图，正式提交时建议结合实物补充电源方案、线材固定方式和接线安全说明。"
        "在当前文档中，本文按通用树莓派加双轴总线舵机云台加手机夹持结构进行说明。"
    )

    add_subtitle(doc, "6. 自主知识产权技术说明")
    add_paragraph(
        doc,
        "从当前仓库源码可以确认，本作品的多端协同架构、设备会话管理、实时推流与预览链路、模板构图与姿态辅助、手势抓拍状态机、云台抽象控制层、后端业务模型以及 AI 任务编排机制均为项目自研集成实现。"
        "项目在实现过程中使用了 FastAPI、SQLAlchemy、Flutter、Vue 3、OpenCV、MediaPipe、pyserial 等开源基础组件，这些组件属于通用框架，并不改变本作品在系统设计、模块集成、交互流程和业务机制上的自主实现属性。"
    )
    add_paragraph(
        doc,
        "经对当前项目源码、目录结构和依赖检查，未发现企业命题专用 SDK、企业设备接口或企业平台强绑定逻辑，因此本文按非企业命题作品填写。"
        "如后续报名信息存在企业命题要求，可在此处补充企业设备、SDK 或协议接入说明。"
    )

    add_heading(doc, "方案实现")
    add_subtitle(doc, "1. 移动端实现")
    add_paragraph(
        doc,
        "移动端是用户主入口，当前代码已经形成较完整的软件模块结构，覆盖认证与配置、首页与导航、本机拍摄、设备联动和历史记录等功能。"
        "在实现上，移动端通过独立业务服务调用云端后端，通过设备服务调用边缘运行时，并保留 WebRTC 联动能力。这样的拆分使移动端同时支持本机独立拍摄和设备联动拍摄两条业务链路。"
    )

    add_subtitle(doc, "2. 边缘设备运行时实现")
    add_paragraph(
        doc,
        "设备运行时是本作品的核心执行平面，负责设备会话创建与关闭、接收移动端图像帧流、执行人体、手势、面部等视觉检测、绘制预览叠加层、控制云台、执行模板构图反馈并完成抓拍。"
        "其中，会话管理器维持单活动会话模型，帧处理器负责图像帧处理与视觉结果输出，运行状态对象存储最新帧、稳定检测、AI 锁定状态与抓拍状态等运行信息，模式管理器负责跟拍、构图、背景锁定等模式组织，抓拍服务负责倒计时与本地文件保存。"
    )

    add_subtitle(doc, "3. 云台控制实现")
    add_paragraph(
        doc,
        "云台控制部分采用抽象驱动加控制器的设计。统一驱动接口用于屏蔽硬件差异，模拟驱动用于无硬件调试，TTL 总线串口驱动用于真实总线舵机控制，上层控制器负责回中、绝对移动、相对移动、反馈刷新和平滑运动。"
        "这样的设计使控制逻辑与具体硬件驱动分离，便于后续替换硬件，同时可以在没有实物的情况下完成大部分软件联调。"
    )

    add_subtitle(doc, "4. 模板构图与手势抓拍实现")
    add_paragraph(
        doc,
        "模板构图与手势抓拍是本作品区别于普通拍照系统的重要能力。模板通过人体关键点、目标框和锚点信息描述理想构图状态；边缘端根据当前人体检测结果与模板目标进行比较，生成构图偏差反馈；"
        "当主体位置和姿态达到要求后，系统进入 ready 状态；手势状态机在 ready 条件满足时支持触发抓拍或强制抓拍。手势抓拍倒计时固定为 3 秒，这一点在会话管理常量定义中可以直接确认。"
    )

    add_subtitle(doc, "5. 云端业务实现")
    add_paragraph(
        doc,
        "后端采用路由层、服务层和模型层的结构。移动端接口已实现登录、注册、当前用户查询、套餐查询、模板管理、会话创建、图片上传、AI 分析、批量选优和历史记录查询等功能。"
        "数据库层当前定义了用户、套餐、订阅、设备、模板、拍摄会话、抓拍记录、AI 任务和 AI Provider 配置等核心数据表，说明系统在数据建模方面已经形成较完整的业务闭环。"
    )
    add_paragraph(
        doc,
        "AI 任务方面，后端将图片分析能力抽象为可配置 Provider，当前实现支持基于 OpenAI-Compatible 接口风格的图像分析调用。"
        "这种设计便于在不改动主业务流程的前提下切换不同模型或不同厂商的服务。"
    )

    add_subtitle(doc, "6. 管理后台实现")
    add_paragraph(
        doc,
        "管理后台提供了与移动端相独立的运营与维护平面。当前源码中可见的典型页面包括工作台、用户管理、推荐模板管理和 AI Provider 配置管理等。"
        "后台的作用不是参与实时拍照执行，而是负责用户、模板、设备和 AI 配置等静态或半静态业务管理，这种职责划分与物联网系统前台实时、后台运维的常见设计模式一致。"
    )

    add_subtitle(doc, "7. 关键技术难点")
    add_bullets(doc, [
        "多端协同复杂度高：系统同时涉及 Flutter 移动端、FastAPI 设备端、FastAPI 云端、Vue 后台和 PostgreSQL 数据库，接口边界与状态同步复杂度明显高于单体应用。",
        "实时链路与业务链路并存：移动端既要访问业务后端，也要与设备运行时保持低时延交互，需要同时处理 REST 与 WebSocket 等多种通信方式。",
        "视觉结果到设备动作的闭环转换：人体检测、姿态判断和手势识别的结果不仅用于显示，还要进一步转化为云台动作和抓拍时机判断，对时序与状态设计要求较高。",
        "模板构图的可执行化：需要把抽象的好看构图转换为模板目标、偏差提示、稳定判断和 ready 状态。",
        "AI 能力的分层使用：设备侧更关注即时辅助，云端侧更关注图片分析和历史沉淀，两者目标不同，需要通过架构分层避免互相拖累。"
    ])

    add_heading(doc, "测试报告")
    add_subtitle(doc, "1. 已有自动化测试")
    table = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["测试类型", "已验证内容", "说明"]):
        table.cell(0, i).text = h
    add_rows(table, [
        ["设备端叠加层与手势测试", "预览叠加层绘制、JPEG 编码、手势抓拍状态机、配置模型、检测配置变更后的重建流程", "证明边缘端核心交互逻辑已具备可测试实现"],
        ["后端抓拍类型测试", "设备联动抓拍类型可持久化", "证明设备联动抓拍已经在后端数据模型中得到支持"],
        ["移动端基础测试", "页面启动与基础界面行为", "说明移动端具备基础测试工程结构"],
    ])
    style_table(table)

    add_subtitle(doc, "2. 各功能模块符合度")
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["功能模块", "设计目标", "当前代码实现情况", "结论"]):
        table.cell(0, i).text = h
    add_rows(table, [
        ["用户登录注册", "完成基础用户认证", "移动端业务接口已提供登录和注册能力", "已实现"],
        ["模板管理", "支持模板创建、查询、删除", "移动端业务接口已提供模板管理能力", "已实现"],
        ["本机图片上传", "支持抓拍文件上传与记录入库", "后端已提供文件上传和抓拍记录入库能力", "已实现"],
        ["设备联动会话", "支持会话打开、关闭、状态查询", "设备运行时已实现会话控制与状态查询接口", "已实现"],
        ["手机推流", "支持移动端向设备端发送图像帧", "设备运行时已实现 WebSocket 推流与帧接口", "已实现"],
        ["设备预览回传", "支持设备端 JPEG 预览输出", "设备运行时已实现图片预览和预览 WebSocket 接口", "已实现"],
        ["云台控制", "支持手动移动、回中、模式切换", "设备控制接口已实现相关能力", "已实现"],
        ["设备抓拍", "支持倒计时触发与文件保存", "设备抓拍接口已实现触发、列表与文件访问能力", "已实现"],
        ["AI 图片分析", "支持后端 AI 分析任务", "后端已实现单图分析、背景分析和批量选优接口", "已实现"],
    ])
    style_table(table)

    add_subtitle(doc, "3. 源码可确认的性能与配置值")
    table = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["指标项", "源码可确认值", "依据", "说明"]):
        table.cell(0, i).text = h
    add_rows(table, [
        ["手势抓拍倒计时", "3.0 秒", "会话管理常量", "属于实际执行逻辑，不是文档假设"],
        ["树莓派性能档检测帧率", "performance 6 FPS / balanced 8 FPS / quality 10 FPS", "树莓派性能档默认配置", "属于默认配置值，不等同于最终实测值"],
        ["树莓派性能档预览帧率", "performance 20 FPS / balanced 24 FPS / quality 24 FPS", "树莓派性能档默认配置", "属于默认配置值"],
        ["预览 JPEG 质量", "76 / 82 / 88 三档", "树莓派性能档默认配置", "用于权衡清晰度与传输开销"],
        ["TTL 串口默认参数", "/dev/ttyUSB0，115200 波特率", "TTL 总线串口默认配置", "说明系统预设了真实硬件串口通信条件"],
        ["网络联调建议", "手机与树莓派同一 5GHz Wi‑Fi 或树莓派热点直连", "项目部署与演示文档", "属于源码配套文档给出的联调建议"],
    ])
    style_table(table)
    add_paragraph(
        doc,
        "需要说明的是，当前仓库中没有完整的压测脚本和统一实测报告，因此本文在测试部分主要填写源码已实现功能和源码可确认配置值。"
        "如果后续补充现场实测帧率、延时和稳定性数据，可进一步增强比赛评审中的说服力。"
    )

    add_heading(doc, "应用前景")
    add_paragraph(
        doc,
        "本作品具有明确的应用场景和进一步扩展潜力。在用户应用层面，可服务于自拍、情侣出游、家庭合照、校园社交、景区打卡和轻量化内容创作等场景。"
        "通过模板构图、设备联动和手势抓拍，系统可以降低用户在拍照过程中的经验门槛和人力依赖。"
    )
    add_paragraph(
        doc,
        "在技术延展层面，当前架构已经具备平台化基础：一方面可以继续接入更多 AI Provider，增强图像分析、构图建议和场景理解能力；"
        "另一方面可以扩展更多云台模式或更多传感器输入，提升物联网设备联动能力。未来可从个人拍照辅助延伸到固定机位内容创作、景区智能拍照点、小型活动自助拍照装置等场景。"
        "若用于更大规模商用，还需继续加强迁移管理、测试覆盖、异常恢复、硬件标准化和安全性设计。"
    )

    add_heading(doc, "参考文献")
    for ref in [
        "[1] 云影随行项目源码仓库：mobile_client、device_runtime、backend、admin_web、database 模块。",
        "[2] FastAPI Official Documentation.",
        "[3] SQLAlchemy Official Documentation.",
        "[4] Flutter Official Documentation.",
        "[5] Vue.js Official Documentation.",
        "[6] OpenCV Official Documentation.",
        "[7] MediaPipe Official Documentation.",
        "[8] PostgreSQL Official Documentation.",
    ]:
        add_paragraph(doc, ref, first_line_chars=None, size=11.5, space_after=3)

    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build_doc()
