from __future__ import annotations

from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables" / "project_research_report"
DOCX_PATH = OUT_DIR / "云影随行-项目研究报告.docx"
ARCH_IMG = OUT_DIR / "research_architecture.png"
FLOW_IMG = OUT_DIR / "research_flow.png"

INK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0x00)
LIGHT_BLUE = "FFFFFF"
LIGHT_GRAY = "FFFFFF"
GRID = "000000"
SONG = "宋体"
HEI = "黑体"
KAI = "楷体"


def ensure_dir() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_run_font(
    run,
    name: str = SONG,
    size: float | None = None,
    *,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = GRID, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def keep_table_row_together(row, *, repeat_header: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if repeat_header and tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def style_table(table, widths_cm: list[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    if widths_cm:
        set_table_width(table, widths_cm)
    for row_idx, row in enumerate(table.rows):
        keep_table_row_together(row, repeat_header=(row_idx == 0))
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, "FFFFFF")
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.15
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_run_font(run, size=9.5, bold=(row_idx == 0), color=INK)


def set_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = SONG
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), SONG)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    for style_name, size, color in (
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 11.5, INK),
    ):
        style = doc.styles[style_name]
        style.font.name = HEI
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEI)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def paragraph_border_bottom(paragraph, color="000000", size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_para(
    doc: Document,
    text: str,
    *,
    size: float = 10.5,
    bold: bool = False,
    font: str = SONG,
    color: RGBColor = INK,
    first_line: bool = True,
    before: float = 0,
    after: float = 5,
    line_spacing: float = 1.25,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, font, size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, HEI, 15 if level == 1 else 12.5, bold=True, color=BLUE)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        bullet = p.add_run("• ")
        set_run_font(bullet, "Calibri", 10.5, bold=True, color=BLUE)
        txt = p.add_run(item)
        set_run_font(txt, SONG, 10.5, color=INK)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, KAI, 9.5, color=MUTED)


def load_font(size: int, *, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        if path and Path(path).exists():
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                pass
    return ImageFont.load_default()


def wrap_zh(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    tokens: list[str] = []
    current_token = ""
    for ch in text:
        if ord(ch) < 128 and not ch.isspace():
            current_token += ch
            continue
        if current_token:
            tokens.append(current_token)
            current_token = ""
        if ch.isspace():
            tokens.append(" ")
        else:
            tokens.append(ch)
    if current_token:
        tokens.append(current_token)

    lines: list[str] = []
    current = ""
    for token in tokens:
        test = current + token
        if draw.textlength(test, font=font) <= max_width:
            current = test
        else:
            if current:
                lines.append(current.rstrip())
            if draw.textlength(token, font=font) <= max_width:
                current = token.lstrip()
            else:
                current = ""
                for ch in token:
                    test = current + ch
                    if draw.textlength(test, font=font) <= max_width:
                        current = test
                    else:
                        if current:
                            lines.append(current.rstrip())
                        current = ch
    if current:
        lines.append(current.rstrip())
    return lines


def box(draw, xy, title: str, body: str, fill="#FFFFFF", outline="#000000") -> None:
    title_font = load_font(25, bold=True)
    body_font = load_font(18)
    x1, y1, x2, y2 = xy
    body = body.rstrip("。；;.")
    draw.rectangle(xy, fill=fill, outline=outline, width=3)
    draw.text((x1 + 18, y1 + 14), title, font=title_font, fill="#000000")
    y = y1 + 56
    max_lines = max(1, int((y2 - y - 14) / 27))
    for line in wrap_zh(draw, body, body_font, x2 - x1 - 36)[:max_lines]:
        draw.text((x1 + 18, y), line, font=body_font, fill="#000000")
        y += 27


def arrow(draw, start, end, color="#000000", width=5) -> None:
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        draw.polygon([(ex, ey), (ex - 16 * sign, ey - 9), (ex - 16 * sign, ey + 9)], fill=color)
    else:
        sign = 1 if ey > sy else -1
        draw.polygon([(ex, ey), (ex - 9, ey - 16 * sign), (ex + 9, ey - 16 * sign)], fill=color)


def generate_diagrams() -> None:
    ensure_dir()
    title_font = load_font(34, bold=True)
    note_font = load_font(20)
    label_font = load_font(17)

    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((80, 55), "系统协同架构图", font=title_font, fill="#000000")

    box(draw, (80, 190, 420, 390), "移动端 App", "手机本地预览、拍照录像、模板管理、设备联动主控、AI结果展示")
    box(draw, (640, 190, 980, 390), "设备运行时", "接收手机帧、OpenCV/MediaPipe检测、云台控制、状态反馈")
    box(draw, (1200, 190, 1540, 390), "业务后端", "账号、套餐、模板、拍摄记录、AI任务、Provider配置")
    box(draw, (640, 560, 980, 760), "云台与硬件", "树莓派或本地主机、TTL总线舵机、双轴云台、mock/真实驱动")
    box(draw, (1200, 560, 1540, 760), "管理后台", "用户、设备、套餐、推荐模板、AI Provider与统计管理")

    arrow(draw, (420, 290), (640, 290))
    draw.text((470, 255), "WebSocket/设备控制", font=label_font, fill="#000000")
    draw.line([(250, 190), (250, 145), (1370, 145)], fill="#000000", width=5)
    arrow(draw, (1370, 145), (1370, 190))
    draw.text((735, 110), "REST/业务与AI", font=label_font, fill="#000000")
    arrow(draw, (810, 390), (810, 560))
    draw.text((835, 470), "串口/控制命令", font=label_font, fill="#000000")
    arrow(draw, (1370, 560), (1370, 390))
    draw.text((1395, 470), "REST/配置治理", font=label_font, fill="#000000")
    img = img.convert("1").convert("RGB")
    img.save(ARCH_IMG)

    img2 = Image.new("RGB", (1600, 920), "#FFFFFF")
    draw2 = ImageDraw.Draw(img2)
    draw2.text((80, 55), "核心业务流程图", font=title_font, fill="#000000")
    steps = [
        ((80, 200, 335, 355), "1. 模板准备", "选择或上传模板，形成结构化template_data"),
        ((420, 200, 675, 355), "2. 进入联动", "手机打开本地CameraPreview，并向设备端建立会话"),
        ((760, 200, 1015, 355), "3. 推送帧流", "手机通过WebSocket将摄像头帧发送至设备运行时"),
        ((1100, 200, 1355, 355), "4. 视觉处理", "设备端检测人体、手势与构图偏差，输出状态"),
        ((1100, 565, 1355, 720), "5. 云台控制", "根据状态发送pan/tilt控制，驱动云台转动"),
        ((760, 565, 1015, 720), "6. 拍摄保存", "最终照片/视频优先由手机保存，并按设置上传后端"),
        ((420, 565, 675, 720), "7. AI分析", "后端按Provider配置生成评分、建议与推荐框"),
    ]
    for xy, title, body in steps:
        box(draw2, xy, title, body, "#FFFFFF")
    arrow(draw2, (335, 278), (420, 278))
    arrow(draw2, (675, 278), (760, 278))
    arrow(draw2, (1015, 278), (1100, 278))
    arrow(draw2, (1228, 355), (1228, 565))
    arrow(draw2, (1100, 642), (1015, 642))
    arrow(draw2, (760, 642), (675, 642))
    arrow(draw2, (548, 565), (548, 355))
    img2 = img2.convert("1").convert("RGB")
    img2.save(FLOW_IMG)


def add_title_page(doc: Document) -> None:
    add_para(doc, "项目研究报告", size=24, bold=True, font=HEI, first_line=False, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "云影随行：基于边缘视觉与大模型协同的智能拍摄系统",
        size=15,
        bold=True,
        font=HEI,
        color=BLUE,
        first_line=False,
        after=24,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run("中国机器人及人工智能创新大赛参赛项目报告")
    set_run_font(run, KAI, 11, color=MUTED)
    paragraph_border_bottom(p)

    rows = [
        ("项目名称", "云影随行：基于边缘视觉与大模型协同的智能拍摄系统"),
        ("项目定位", "面向大众手机摄影、轻量化自助拍摄和内容创作辅助的物联网智能拍摄系统"),
        ("系统组成", "移动端 App、设备运行时、双轴云台、业务后端、管理后台与数据库"),
        ("核心能力", "手机主预览、模板构图、设备联动、云台跟随、手机拍摄保存、后端 AI Provider 分析"),
        ("完成日期", "2026 年 5 月 26 日"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "内容"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    style_table(table, [3.5, 12.2])
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    entries = [
        "一、摘要",
        "二、研究背景与问题定义",
        "三、项目目标与总体方案",
        "四、系统总体设计",
        "五、关键技术路线",
        "六、功能实现与业务流程",
        "七、接口、数据与安全设计",
        "八、测试验证与风险分析",
        "九、创新性、实用价值与应用前景",
        "十、结论与后续研究方向",
        "参考文献",
    ]
    for entry in entries:
        add_para(doc, entry, first_line=False, after=3)
    doc.add_page_break()


def add_report_body(doc: Document) -> None:
    add_heading(doc, "一、摘要", 1)
    add_para(
        doc,
        "“云影随行”是一套面向大众手机摄影、校园活动记录、旅行打卡和轻量化内容创作场景的物联网智能拍摄系统。系统以手机 App 为用户入口，以树莓派或本机设备运行时为边缘处理节点，以双轴云台/TTL 总线舵机为执行机构，以业务后端和管理后台为数据与运营支撑，形成“手机负责画面、设备负责转动、服务器负责智能、后台负责管理”的多端协同架构。",
    )
    add_para(
        doc,
        "项目围绕“手机负责画面、设备负责转动、服务器负责智能、后台负责管理”的设计原则展开。设备联动场景下，主画面与最终媒体均来自手机本地 CameraPreview 和相册保存；设备运行时接收手机推送的低延迟帧，用于人体/手势检测、模板构图评分、目标跟踪和云台控制；后端统一承担账号、套餐、模板、媒体记录、AI 任务和 Provider 配置管理。",
    )
    add_para(
        doc,
        "研究结果表明，项目的核心价值不在于发明 WebSocket、人体姿态检测或视觉大模型本身，而在于将移动端拍摄、边缘视觉处理、云台执行、结构化模板构图、后端 AI Provider 治理和管理后台沉淀组织成可演示、可维护、可扩展的完整工程闭环。",
    )
    add_para(doc, "关键词：智能拍摄；边缘视觉；模板构图；手机云台；AI Provider；物联网应用", first_line=False, bold=True)

    add_heading(doc, "二、研究背景与问题定义", 1)
    add_heading(doc, "2.1 研究背景", 2)
    add_para(
        doc,
        "手机摄影已经从单纯记录工具发展为社交分享、校园活动、旅行留念和短视频创作的高频入口。普通用户在真实拍摄中经常同时面对取景、站位、角度调整、抓拍时机和后期筛选等问题。传统自拍杆解决的是距离问题，手机稳定器解决的是防抖和基础跟踪问题，拍照/修图 App 更偏向后期处理；这些方案很少把拍前模板准备、拍中设备联动和拍后 AI 分析整合为统一闭环。",
    )
    add_heading(doc, "2.2 目标用户与核心痛点", 2)
    add_bullets(
        doc,
        [
            "自拍、合影、旅行打卡用户需要减少对他人帮拍的依赖，并获得更稳定的构图结果。",
            "短视频与轻量内容创作者需要在固定机位或云台辅助下快速完成角度探索、跟随和结果筛选。",
            "校园活动、文旅自助拍摄等场景需要把设备状态、拍摄记录、AI 分析结果和后台配置统一沉淀，便于复用和运营。",
            "开发和教学场景需要可在无硬件条件下演示主要流程，同时保留真实舵机联调能力。",
        ],
    )
    add_heading(doc, "2.3 现有方案对比", 2)
    table = doc.add_table(rows=1, cols=5)
    for idx, h in enumerate(["维度", "传统自拍杆", "手机稳定器", "拍照/修图 App", "云影随行"]):
        table.cell(0, idx).text = h
    for row in [
        ["主画面与最终媒体", "手机本地", "手机本地", "手机本地", "手机本地 CameraPreview 与相册保存"],
        ["设备联动", "手动调节", "厂商闭环云台", "无硬件联动", "手机主控 + 设备运行时 + TTL 云台"],
        ["构图辅助", "依赖经验", "部分跟踪", "偏后期建议", "模板数据、人体关键点、结构化 overlay"],
        ["AI 分析", "无", "部分机型内置", "偏后期处理", "后端 Provider 统一调用与任务记录"],
        ["可扩展性", "低", "受厂商限制", "软件内闭环", "移动端、设备端、后端、后台分层扩展"],
    ]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, [2.8, 3.0, 3.0, 3.0, 4.3])
    add_caption(doc, "表 1  现有方案与本项目对比")

    add_heading(doc, "三、项目目标与总体方案", 1)
    add_para(
        doc,
        "项目目标是构建一套可演示、可维护、可扩展的智能拍摄系统，使普通用户在自拍、合影、旅行打卡、校园活动记录和轻量内容创作中，能够更稳定地完成取景、构图、跟随、拍摄保存和结果分析。系统不把某一个端作为万能中心，而是让移动端、设备端、业务后端和管理后台分别承担最适合自己的职责。",
    )
    table = doc.add_table(rows=1, cols=3)
    for idx, h in enumerate(["设计目标", "实现方式", "项目价值"]):
        table.cell(0, idx).text = h
    rows = [
        ["主画面可信", "手机本地 CameraPreview 作为设备联动主画面，最终照片/视频优先保存到手机相册", "用户看到的画面与保存结果一致，降低设备回传画面带来的复杂度"],
        ["设备可控", "设备运行时接收手机帧，执行 OpenCV/MediaPipe 检测、模板评分和 pan/tilt 控制", "让普通云台具备自动跟随、构图辅助和手势触发能力"],
        ["智能可治理", "后端统一维护 AI Provider、AI 任务、套餐能力和结构化响应", "避免在手机或设备端保存生产 AI Key，便于模型替换和成本管理"],
        ["演示可落地", "mock 与 ttl_bus 驱动并存，支持本机联调和树莓派真实舵机运行", "兼顾课堂开发、竞赛演示和硬件现场验证"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, [3.6, 6.0, 5.8])
    add_caption(doc, "表 2  项目目标与实现方式")

    add_heading(doc, "四、系统总体设计", 1)
    add_para(
        doc,
        "项目采用四端协同架构，将实时控制、最终媒体、业务数据和后台运营分别放在最适合的位置。移动端掌握用户主画面和最终照片/视频，保证“用户看到什么就保存什么”；设备端运行在树莓派或本机环境中，专注低时延检测、跟踪和云台控制；业务后端统一管理账号、套餐、模板、媒体记录和 AI Provider；管理后台负责运营与配置。",
    )
    doc.add_picture(str(ARCH_IMG), width=Inches(6.35))
    add_caption(doc, "图 1  系统协同架构图")
    table = doc.add_table(rows=1, cols=3)
    for idx, h in enumerate(["模块", "主要技术", "当前职责"]):
        table.cell(0, idx).text = h
    rows = [
        ["mobile_client", "Flutter、camera、ML Kit、WebSocket、SharedPreferences", "手机本地预览、拍照录像、相册保存、模板构图、设备联动主控、AI 结果展示"],
        ["device_runtime", "FastAPI、OpenCV、MediaPipe、pyserial、WebSocket、aiortc", "接收手机帧、本地检测、目标跟踪、模板评分、手势状态、TTL 云台控制"],
        ["backend", "FastAPI、SQLAlchemy、PostgreSQL、httpx、Pydantic", "用户、套餐、订阅、设备、模板、媒体记录、AI 任务和 Provider 调用"],
        ["admin_web", "Vue 3、Vite、Element Plus、Pinia、Axios", "用户、套餐、设备、推荐模板、媒体、AI 任务和 Provider 配置管理"],
        ["database", "PostgreSQL SQL", "数据表、约束、索引、触发器和兼容旧库的结构说明"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, [3.0, 5.2, 7.3])
    add_caption(doc, "表 3  系统模块与职责")

    add_heading(doc, "五、关键技术路线", 1)
    add_heading(doc, "5.1 手机主画面与帧上行", 2)
    add_para(
        doc,
        "当前实现中，拍摄页和设备联动页都以手机本地相机作为主画面。设备联动时，手机端通过 CameraController 初始化 CameraPreview，并通过 ImageStream 抽取帧；主推流路径为 WebSocket /api/device/stream/mobile-ws，设备端当前支持 nv21 格式，并使用 OpenCV 的 COLOR_YUV2BGR_NV21 转为 BGR 图像。HTTP JPEG 帧上传作为兼容兜底，设备端调试预览与 WebRTC signaling 保留，但不作为用户主预览的生产依赖。",
    )
    add_heading(doc, "5.2 边缘视觉与云台控制", 2)
    add_para(
        doc,
        "device_runtime 的核心职责是“让设备动起来”：接收手机帧，运行 OpenCV/MediaPipe 检测，生成稳定目标、模板构图得分、手势状态和跟踪命令，再由 GimbalController 控制 pan/tilt 两轴。开发环境可使用 mock 驱动，无需真实硬件即可完成联调；树莓派现场运行时可切换到 ttl_bus，通过串口控制 TTL 总线舵机。",
    )
    add_heading(doc, "5.3 模板构图与归一化坐标", 2)
    add_para(
        doc,
        "模板构图不是单独训练的构图模型，而是围绕 template_data 的工程化结构设计。模板数据包含 bbox、head_box、points、segments、shoulder_center、face_center 等归一化字段；移动端负责识别和绘制，设备端 TemplateComposeEngine 根据实时人物点与模板点的对齐误差、姿态相似度、完整度、人物面积比例和朝向等加权项计算评分，源码中的 ready 阈值为 68.0。所有跨端视觉坐标均采用 0..1 归一化坐标，手机端再根据预览尺寸和裁剪方式转换为屏幕坐标。",
    )
    add_heading(doc, "5.4 后端 AI Provider 治理", 2)
    add_para(
        doc,
        "AI 能力统一由 backend 管理，手机端与设备端不保存生产 AI Key。后端通过 ai_provider_configs 维护 OpenAI-compatible Provider，并结合套餐 feature_flags 控制可用能力。照片分析、背景分析、连拍选优、多角度扫描等任务由 MobileService 创建并写入 ai_tasks，AI 响应被规范为 score、summary、best_candidate_index、recommended_pan_delta、recommended_tilt_delta 和 target_box_norm 等结构化字段。",
    )

    add_heading(doc, "六、功能实现与业务流程", 1)
    doc.add_picture(str(FLOW_IMG), width=Inches(6.35))
    add_caption(doc, "图 2  核心业务流程图")
    doc.add_page_break()
    table = doc.add_table(rows=1, cols=3)
    for idx, h in enumerate(["流程", "当前实现", "说明"]):
        table.cell(0, idx).text = h
    rows = [
        ["手机独立拍摄", "普通拍照、录像、模板构图、背景分析、AI 连拍选优、历史记录", "最终媒体优先保存到手机相册，按设置上传后端"],
        ["设备联动基础", "健康检查、会话打开、手机推流、手动控制、回中、灵敏度调整", "主画面仍来自手机本地相机"],
        ["自动跟随", "选择肩部/人脸中心点，设备端根据目标偏移输出 pan/tilt 命令", "使用归一化坐标与跟踪控制器限幅"],
        ["模板构图", "选择模板，按肩部或面部中心点对齐到模板点", "TemplateComposeEngine 输出评分与 ready 状态"],
        ["AI 自动找角度", "手机控制云台扫描候选角度并缓存候选帧，调用 analyze-scan", "后端返回最佳候选、分数、原因，手机转回最佳角度"],
        ["背景锁定", "多角度扫描后返回最佳背景与 target_box_norm", "手机端绘制 AI 推荐框，可显示/隐藏"],
        ["手势抓拍", "启用手部 landmarks，检测后进入 3 秒倒计时与冷却", "降低瞬时识别抖动导致的重复触发"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, [3.2, 7.0, 5.2])
    add_caption(doc, "表 4  主要功能流程")

    add_heading(doc, "七、接口、数据与安全设计", 1)
    add_heading(doc, "7.1 核心接口", 2)
    table = doc.add_table(rows=1, cols=3)
    for idx, h in enumerate(["接口层", "典型路径", "作用"]):
        table.cell(0, idx).text = h
    rows = [
        ["手机到后端", "/api/mobile/auth、/templates、/captures、/ai/*", "登录注册、模板、会话、上传、AI 分析与历史记录"],
        ["管理后台到后端", "/api/admin/users、/plans、/devices、/captures、/ai/tasks、/ai/provider-configs", "运营管理与 Provider 配置"],
        ["手机到设备端", "/api/device/session、/control、/status、/stream/mobile-ws", "会话、云台控制、状态读取和手机帧上行"],
        ["设备调试/扩展", "/api/device/preview-ws、/preview.jpg、/webrtc/offer", "调试预览和后续实时媒体实验"],
        ["多角度扫描", "POST /api/mobile/ai/analyze-scan", "上传 1 到 12 张候选图及角度元数据，返回最佳候选与结构化结果"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, [3.2, 5.5, 6.8])
    add_caption(doc, "表 5  核心接口概览")

    add_heading(doc, "7.2 数据契约", 2)
    add_bullets(
        doc,
        [
            "视觉坐标统一使用 0..1 归一化坐标，避免不同分辨率、横竖屏和裁剪方式导致跨端错位。",
            "template_data 保存人物框、头部框、关键点、骨架线、肩部中心、面部中心和图片来源，用于模板复拍和设备端对齐。",
            "target_box_norm 推荐使用对象格式 x、y、w、h、label；后端兼容旧数组格式，但新接口应使用对象。",
            "AI 任务统一写入 ai_tasks，便于移动端展示、后台追踪和失败降级记录。",
        ],
    )
    add_heading(doc, "7.3 安全与隐私", 2)
    add_para(
        doc,
        "项目文档明确要求不提交 .env、真实数据库密码、认证密钥、AI Provider API Key、用户照片/视频、日志、本地数据库、模型权重、构建产物等敏感或可再生成资源。生产 AI Key 只应保存在后端 Provider 配置中，手机端与树莓派端不直接保存。设备联调时真机必须使用局域网 IP，不能把手机自身的 127.0.0.1 当作后端地址。",
    )

    add_heading(doc, "八、测试验证与风险分析", 1)
    add_para(
        doc,
        "本报告编写过程主要进行了文档与源码层面的核验，没有替代真实树莓派、舵机、供电和现场网络的长时间稳定性测试。根据项目文档与现有测试入口，后续正式提交或路演前应重点完成以下验证。",
    )
    table = doc.add_table(rows=1, cols=3)
    for idx, h in enumerate(["验证项", "建议方法", "通过标准"]):
        table.cell(0, idx).text = h
    rows = [
        ["移动端基础", "flutter analyze、flutter test、真机拍照录像", "无静态错误；照片/视频进入手机相册"],
        ["后端基础", "启动 PostgreSQL 与 backend，访问 /api/health", "健康检查通过，登录、模板、上传接口可用"],
        ["设备端基础", "mock 驱动启动 device_runtime，访问 /api/device/health", "设备状态正常，手动控制接口可响应"],
        ["手机推流", "设备联动页开启手机推流，观察 /api/device/status", "设备端能接收帧并返回检测/跟踪状态"],
        ["真实云台", "树莓派上设置 DEVICE_SERVO_DRIVER=ttl_bus 并接入舵机", "pan/tilt 移动、回中、限幅和供电稳定"],
        ["AI 链路", "后台配置 Provider，触发 analyze-photo/analyze-scan", "AI 任务写入，移动端展示评分、摘要和推荐框"],
        ["现场稳定性", "5GHz Wi-Fi 或设备热点下连续演示", "无明显卡顿、重复触发或云台异常越界"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, [3.2, 6.8, 5.4])
    add_caption(doc, "表 6  建议验证清单")
    add_heading(doc, "8.1 主要风险", 2)
    add_bullets(
        doc,
        [
            "网络风险：局域网质量会影响手机帧上行和控制响应，应优先使用 5GHz Wi-Fi 或设备热点。",
            "硬件风险：TTL 舵机供电、共地、负载和线缆稳定性需要单独测试，不能只凭 mock 驱动判断真实效果。",
            "AI 风险：Provider 可用性、费用、超时和返回 JSON 稳定性会影响自动找角度与背景锁定，应保留失败降级提示。",
            "性能风险：完整人体、手部、人脸检测同时开启会增加边缘端负载，现场演示可按需关闭非必要 overlay。",
            "表述风险：不能把成熟开源技术、第三方视觉模型或未实测指标写成项目自研成果。",
        ],
    )

    add_heading(doc, "九、创新性、实用价值与应用前景", 1)
    add_heading(doc, "9.1 创新性分析", 2)
    add_bullets(
        doc,
        [
            "架构创新：把手机端、设备运行时、业务后端和管理后台组合成面向智能拍摄的完整物联网应用，而非单点拍照工具。",
            "职责边界创新：明确手机负责主画面和最终媒体，设备负责转动和状态闭环，服务器负责 AI 与业务治理。",
            "模板数据契约创新：把模板从静态图片提升为可跨端复用的结构化数据，支持显示、对齐、评分和历史沉淀。",
            "AI 治理创新：通过后端 Provider 配置和套餐 feature_flags 统一管理视觉模型能力，便于替换、分级、追踪和成本控制。",
            "工程可演示性创新：mock 与 ttl_bus 驱动并存，让无硬件开发、课堂演示和真实舵机联调可以共用同一套接口。",
        ],
    )
    add_heading(doc, "9.2 实用价值", 2)
    add_para(
        doc,
        "对普通用户而言，系统能够通过模板构图、手势倒计时、自动跟随和 AI 候选分析减少反复试拍；对教学和竞赛而言，系统展示了移动端、边缘端、后端和后台的完整协同；对后续产品化而言，后台 Provider 配置、套餐能力控制和媒体/AI 任务记录为服务运营留下扩展空间。",
    )
    table = doc.add_table(rows=1, cols=3)
    for idx, h in enumerate(["应用场景", "可用能力", "实际价值"]):
        table.cell(0, idx).text = h
    rows = [
        ["个人自拍/合影", "手机主预览、模板构图、手势倒计时", "减少对他人帮拍的依赖，提高构图稳定性"],
        ["旅行记录/校园活动", "设备联动、自动跟随、历史记录", "适合固定机位或轻量云台辅助拍摄"],
        ["短视频/直播辅助", "目标跟随、背景锁定、多角度候选分析", "帮助创作者快速探索合适角度"],
        ["教学与竞赛演示", "mock/ttl_bus 双驱动、REST/WebSocket 接口", "展示物联网软硬协同与多端工程结构"],
        ["文旅自助拍摄", "后台配置、AI 分析、模板管理", "可扩展为自助拍摄终端或模板化服务"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, [3.2, 6.4, 5.8])
    add_caption(doc, "表 7  应用场景与价值")

    add_heading(doc, "十、结论与后续研究方向", 1)
    add_para(
        doc,
        "云影随行围绕手机摄影中“看不准、站不好、调不稳、选不出”的实际痛点，构建了一个由移动端、设备运行时、业务后端、管理后台和数据库组成的物联网智能拍摄系统。项目已经形成清晰工程边界：手机负责主画面与最终媒体，设备负责本地检测与云台转动，服务器负责 AI 与业务数据，后台负责管理配置。该架构强调手机主预览和后端 AI 治理，既保证用户拍摄体验，也降低设备回传画面作为主预览所带来的链路复杂度。",
    )
    add_para(
        doc,
        "后续研究可从四个方向继续推进：一是完善真实硬件外壳、供电安全、舵机负载和长时间稳定性测试；二是扩展多设备管理、蓝牙/热点配网和远程运维能力；三是增强多人构图、语音控制、模板市场和 AI 创作能力；四是建立更系统的用户评测和场景数据集，用实际拍摄效果验证模板构图与多角度 AI 分析对出片质量的提升。",
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "Flutter Team. camera package / Flutter camera cookbook.",
        "Google for Developers. ML Kit Pose Detection.",
        "Google AI Edge. MediaPipe Pose Landmarker.",
        "OpenCV Documentation. Color Space Conversions.",
        "RFC Editor. RFC 6455: The WebSocket Protocol.",
        "FastAPI Documentation.",
        "OpenAI Documentation. Images and vision.",
        "PostgreSQL Documentation.",
        "Vue.js Documentation.",
    ]
    for idx, ref in enumerate(refs, start=1):
        add_para(doc, f"[{idx}] {ref}", first_line=False, after=3)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—  ")
    set_run_font(run, SONG, 9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    footer._p.append(fld)
    run = footer.add_run("  —")
    set_run_font(run, SONG, 9, color=MUTED)


def build() -> None:
    ensure_dir()
    generate_diagrams()
    doc = Document()
    set_style(doc)
    add_footer(doc)
    add_title_page(doc)
    add_toc(doc)
    add_report_body(doc)
    doc.save(str(DOCX_PATH))
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
