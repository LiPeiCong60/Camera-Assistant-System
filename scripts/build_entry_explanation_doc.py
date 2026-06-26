from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables" / "entry_explanation_2026"
DESKTOP = Path.home() / "Desktop"

SIGNED_DOCX = DESKTOP / "参赛作品说明书-署名版.docx"
ANON_DOCX = DESKTOP / "参赛作品说明书-匿名版.docx"

PROJECT_NAME = "云影随行：基于边缘视觉与大模型协同的智能拍摄系统"
SCHOOL = "福州大学至诚学院"
DEPARTMENT = "计算机工程系"
CLASS_NAME = "计算机科学与技术专业 3 班"
STUDENTS = "李培聪、薛宇航、林少钢、李泽烨"
ADVISOR = "李娟"
FINISH_TIME = "2026 年 5 月"

BLACK = RGBColor(0, 0, 0)


def ensure_dir() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_run_font(run, font_name: str = "宋体", size_pt: float = 10.5, *, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = BLACK


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "000000", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def style_table(table) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = 1.12
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_run_font(run, "宋体", 9.5, bold=(row_idx == 0))
        if row_idx == 0:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.45)
    section.bottom_margin = Cm(2.25)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK

    for style_name, font, size in (
        ("Title", "黑体", 22),
        ("Heading 1", "黑体", 15),
        ("Heading 2", "黑体", 12),
        ("Heading 3", "宋体", 11),
    ):
        style = doc.styles[style_name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK


def add_para(
    doc: Document,
    text: str,
    *,
    first_line: bool = True,
    after: float = 4,
    before: float = 0,
    size: float = 10.5,
    font: str = "宋体",
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.first_line_indent = Cm(0.74) if first_line else Cm(0)
    para.paragraph_format.line_spacing = 1.28
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.space_before = Pt(before)
    run = para.add_run(text)
    set_run_font(run, font, size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph(style=f"Heading {level}")
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    para.paragraph_format.space_after = Pt(5 if level == 1 else 3)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_run_font(run, "黑体", 15 if level == 1 else 12, bold=True)


def add_caption(doc: Document, text: str) -> None:
    add_para(doc, text, first_line=False, after=6, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_numbered_items(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, 1):
        add_para(doc, f"（{idx}）{item}", first_line=True, after=3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    if len(headers) == 3:
        widths = [Cm(3.2), Cm(4.9), Cm(7.9)]
    elif len(headers) == 4:
        widths = [Cm(3.4), Cm(4.2), Cm(4.2), Cm(4.2)]
    else:
        widths = [Cm(16 / len(headers)) for _ in headers]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    style_table(table)


def load_font(size: int, *, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            try:
                return ImageFont.truetype(path, size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        trial = current + ch
        if draw.textlength(trial, font=font) <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_box(draw: ImageDraw.ImageDraw, xy, title: str, body: str) -> None:
    x1, y1, x2, y2 = xy
    title_font = load_font(26, bold=True)
    body_font = load_font(21)
    draw.rectangle(xy, fill="white", outline="black", width=3)
    draw.text((x1 + 18, y1 + 14), title, font=title_font, fill="black")
    y = y1 + 54
    for line in wrap_text(draw, body, body_font, x2 - x1 - 36)[:4]:
        draw.text((x1 + 18, y), line, font=body_font, fill="black")
        y += 29


def draw_arrow(draw: ImageDraw.ImageDraw, start, end) -> None:
    draw.line([start, end], fill="black", width=4)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        draw.polygon([(ex, ey), (ex - 16 * sign, ey - 9), (ex - 16 * sign, ey + 9)], fill="black")
    else:
        sign = 1 if ey > sy else -1
        draw.polygon([(ex, ey), (ex - 9, ey - 16 * sign), (ex + 9, ey - 16 * sign)], fill="black")


def make_architecture_diagram() -> Path:
    path = OUT_DIR / "说明书-系统协同架构图.png"
    img = Image.new("RGB", (1500, 820), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(34, bold=True)
    draw.text((48, 34), "系统协同架构图", font=title_font, fill="black")
    draw_box(draw, (70, 160, 390, 340), "移动端 App", "手机本地预览、拍照录像、模板管理、设备联动主控、AI 结果展示")
    draw_box(draw, (590, 160, 910, 340), "设备运行时", "接收手机帧、OpenCV/MediaPipe 检测、云台控制、状态反馈")
    draw_box(draw, (1110, 160, 1430, 340), "业务后端", "账号、套餐、模板、拍摄记录、AI 任务、Provider 配置")
    draw_box(draw, (590, 510, 910, 690), "云台与硬件", "树莓派或本地主机、TTL 总线舵机、双轴云台、mock/真实驱动")
    draw_box(draw, (1110, 510, 1430, 690), "管理后台", "用户、设备、套餐、推荐模板、AI Provider 与统计管理")
    draw_arrow(draw, (390, 250), (590, 250))
    draw.text((420, 220), "WebSocket/REST", font=load_font(20), fill="black")
    draw_arrow(draw, (910, 250), (1110, 250))
    draw.text((945, 220), "REST/AI 任务", font=load_font(20), fill="black")
    draw_arrow(draw, (750, 340), (750, 510))
    draw.text((770, 410), "串口/控制命令", font=load_font(20), fill="black")
    draw_arrow(draw, (1270, 340), (1270, 510))
    draw.text((1290, 410), "配置治理", font=load_font(20), fill="black")
    img.save(path)
    return path


def make_flow_diagram() -> Path:
    path = OUT_DIR / "说明书-核心业务流程图.png"
    img = Image.new("RGB", (1500, 840), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(34, bold=True)
    draw.text((48, 34), "核心业务流程图", font=title_font, fill="black")
    boxes = [
        (70, 170, 310, 320, "1. 模板准备", "上传或选择模板，生成结构化 template_data"),
        (390, 170, 630, 320, "2. 进入联动", "手机打开本地 CameraPreview，并向设备端建立会话"),
        (710, 170, 950, 320, "3. 推送帧流", "手机通过 WebSocket 将摄像头帧发送至设备运行时"),
        (1030, 170, 1270, 320, "4. 视觉处理", "设备端检测人体、手势与构图偏差，输出状态"),
        (390, 520, 630, 670, "5. 云台控制", "设备端或手机端根据模式发送 pan/tilt 控制"),
        (710, 520, 950, 670, "6. 拍摄保存", "最终照片/视频优先由手机保存，并按设置上传后端"),
        (1030, 520, 1270, 670, "7. AI 分析", "后端按 Provider 配置生成评分、建议与推荐框"),
    ]
    for b in boxes:
        draw_box(draw, b[:4], b[4], b[5])
    draw_arrow(draw, (310, 245), (390, 245))
    draw_arrow(draw, (630, 245), (710, 245))
    draw_arrow(draw, (950, 245), (1030, 245))
    draw_arrow(draw, (1150, 320), (1150, 460))
    draw_arrow(draw, (1030, 595), (950, 595))
    draw_arrow(draw, (710, 595), (630, 595))
    draw_arrow(draw, (510, 520), (510, 320))
    img.save(path)
    return path


def add_cover(doc: Document, anonymous: bool) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("参赛作品说明书")
    set_run_font(run, "黑体", 24, bold=True)
    for _ in range(5):
        doc.add_paragraph()

    fields = [
        ("作品名称", PROJECT_NAME),
    ]
    if not anonymous:
        fields.extend(
            [
                ("学校", SCHOOL),
                ("学院（系）", DEPARTMENT),
                ("专业班别", CLASS_NAME),
                ("学生姓名", STUDENTS),
                ("指导教师", ADVISOR),
            ]
        )
    fields.append(("完成时间", FINISH_TIME))

    for label, value in fields:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(3.2 if anonymous else 2.1)
        para.paragraph_format.space_after = Pt(9)
        para.paragraph_format.line_spacing = 1.4
        r1 = para.add_run(f"{label}：")
        set_run_font(r1, "宋体", 12, bold=True)
        r2 = para.add_run(value)
        set_run_font(r2, "宋体", 12)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    contents = [
        "一、作品简介",
        "二、设计原理",
        "三、队员分工",
        "四、创新点",
        "五、实用点",
        "六、总结",
    ]
    for item in contents:
        add_para(doc, item, first_line=False, after=4, size=11)
    doc.add_page_break()


def add_intro(doc: Document) -> None:
    add_heading(doc, "一、作品简介", 1)
    add_para(
        doc,
        "“云影随行”是一套面向大众手机摄影、移动内容创作和轻量化自助拍摄场景的物联网智能拍摄系统。作品以手机 App 为用户入口，以树莓派或本机设备运行时为边缘处理节点，以双轴云台/舵机为执行机构，以业务后端和管理后台为数据与运营支撑，形成“手机负责画面、设备负责转动、服务器负责智能、后台负责管理”的多端协同方案。",
    )
    add_para(
        doc,
        "根据当前项目源码与文档，系统已经覆盖移动端拍摄、模板引导、设备联动、云台控制、人体/手势/构图状态反馈、拍摄记录、AI 图片分析、AI Provider 配置和管理后台等核心能力。与参考技术文档相比，当前实现更明确地强调手机本地 CameraPreview 作为用户主预览与最终媒体来源，设备端主要负责接收手机帧、做必要的视觉处理与云台控制；设备端 JPEG 预览和 WebRTC 路径保留为调试、兼容与后续扩展能力。",
    )
    add_para(
        doc,
        "作品面向的实际问题是：普通用户在自拍、合影、旅行记录、校园活动和短视频创作中，常常需要同时完成取景、站位、角度调整和抓拍，既依赖经验，也依赖他人协助。本作品通过模板构图、自动跟随、手势倒计时、AI 多角度分析和后台记录沉淀，降低普通用户获得稳定构图和优质影像的门槛。",
    )

    add_table(
        doc,
        ["模块", "主要技术", "当前职责"],
        [
            ["mobile_client", "Flutter / camera\nHTTP / WebSocket\nSharedPreferences", "负责手机本地预览、拍照录像、模板管理、历史记录、设备联动主控和 AI 结果展示。"],
            ["device_runtime", "FastAPI / OpenCV\nMediaPipe / WebSocket\npyserial", "负责接收手机帧、运行视觉检测、计算跟踪/构图状态、控制云台并返回设备状态。"],
            ["backend", "FastAPI / SQLAlchemy\nPostgreSQL / httpx", "负责账号、套餐、模板、拍摄记录、AI 任务和 AI Provider 的统一管理。"],
            ["admin_web", "Vue 3 / Pinia\nElement Plus / Axios", "负责用户、设备、套餐、推荐模板、拍摄记录、AI 任务和 Provider 配置维护。"],
            ["database", "PostgreSQL", "负责用户、套餐、模板、会话、媒体记录、AI 任务和 Provider 配置等结构化数据持久化。"],
        ],
    )
    add_caption(doc, "表 1  系统模块与职责")


def add_design(doc: Document, arch_img: Path, flow_img: Path) -> None:
    add_heading(doc, "二、设计原理", 1)
    add_heading(doc, "（一）总体设计思路", 2)
    add_para(
        doc,
        "系统采用“移动端 + 边缘设备端 + 云端业务后端 + 管理后台”的分层协同设计。移动端负责用户交互、手机相机采集、拍摄保存和设备联动主控；设备运行时负责局域网内的低时延状态闭环、视觉检测与云台执行；业务后端负责账户、套餐、模板、媒体记录和 AI 任务沉淀；管理后台负责运营配置和数据治理。",
    )
    add_para(
        doc,
        "这种设计避免把所有能力堆在单一端上：实时控制与硬件动作放在本地设备侧，长期保存、权限控制、套餐分级和 AI Provider 管理放在后端，手机端保留主画面和最终拍摄权，确保用户看到的画面与最终保存的照片/视频一致。",
    )
    doc.add_picture(str(arch_img), width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图 1  系统协同架构图")

    add_heading(doc, "（二）关键链路原理", 2)
    add_para(
        doc,
        "设备联动时，手机端先初始化本地相机并显示本地预览，再通过 WebSocket 将摄像头帧推送给 device_runtime。设备端将帧转换为可处理图像后，进入 OpenCV/MediaPipe 检测、构图评分、手势状态识别和跟踪控制流程，并通过 HTTP 状态接口或调试预览接口把结果返回给手机端。手机端根据归一化坐标绘制人体框、骨架线、模板框和 AI 推荐框。",
    )
    add_para(
        doc,
        "云台控制采用抽象驱动加控制器的设计。开发和演示阶段可使用 mock 驱动完成无硬件联调；真实设备侧可使用 ttl_bus 驱动控制 TTL 总线舵机。控制层限制角度范围、移动速度和命令节奏，避免设备动作过猛或越界。手势抓拍采用 3 秒倒计时和冷却机制，减少瞬时识别抖动导致的误触发。",
    )
    doc.add_picture(str(flow_img), width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图 2  核心业务流程图")
    doc.add_page_break()

    add_para(
        doc,
        "AI 与数据闭环方面，AI 能力由后端统一配置和调用，手机端与设备端不保存生产 AI Key。后端通过 ai_provider_configs 维护 Provider，并可结合套餐 feature_flags 控制不同用户可使用的 AI 能力。当前业务侧支持照片分析、背景分析、连拍优选和多角度候选图分析，AI 输出被规范为评分、摘要、推荐角度、推荐框 target_box_norm 等结构化字段，便于写入 ai_tasks 并在移动端展示。",
    )
    add_para(
        doc,
        "模板构图则通过 template_data 在多端复用。模板数据包含人物框、头部框、关键点、骨架线、肩部中心和面部中心等归一化信息，移动端、后端与设备端都围绕这一结构进行识别、保存、读取与对齐。这样的数据契约比只保存图片更利于后续复拍、推荐和自动构图。",
    )


def add_team(doc: Document, anonymous: bool) -> None:
    add_heading(doc, "三、队员分工", 1)
    add_para(
        doc,
        "本作品按照“移动端体验、设备运行时、云端后端、管理后台与测试文档”进行模块化协作。以下分工根据当前项目模块和成稿需要整理；匿名版以成员编号代替姓名，不出现学校、院系、专业班别、学生姓名和指导教师等识别信息。",
    )
    if anonymous:
        rows = [
            ["成员 A", "总体架构、后端业务、设备运行时主链路", "负责系统架构拆分、核心接口设计、后端业务模型、设备会话与控制链路整理。"],
            ["成员 B", "移动端与设备联动交互", "负责 Flutter 页面、手机本地预览、拍照录像、设备联动控制、模板选择与状态展示。"],
            ["成员 C", "管理后台、数据库与 AI 配置", "负责后台管理页面、套餐与 Provider 配置、数据库结构梳理和运营配置流程。"],
            ["成员 D", "硬件联调、部署测试与材料整理", "负责树莓派/本机运行环境、云台舵机联调、演示流程、测试记录和说明书图表整理。"],
        ]
    else:
        rows = [
            ["李培聪", "总体架构、后端业务、设备运行时主链路", "负责系统架构拆分、核心接口设计、后端业务模型、设备会话与控制链路整理。"],
            ["薛宇航", "移动端与设备联动交互", "负责 Flutter 页面、手机本地预览、拍照录像、设备联动控制、模板选择与状态展示。"],
            ["林少钢", "管理后台、数据库与 AI 配置", "负责后台管理页面、套餐与 Provider 配置、数据库结构梳理和运营配置流程。"],
            ["李泽烨", "硬件联调、部署测试与材料整理", "负责树莓派/本机运行环境、云台舵机联调、演示流程、测试记录和说明书图表整理。"],
        ]
    add_table(doc, ["成员", "主要分工", "具体工作"], rows)
    add_caption(doc, "表 2  队员分工表")


def add_innovation(doc: Document) -> None:
    add_heading(doc, "四、创新点", 1)
    add_numbered_items(
        doc,
        [
            "多端协同架构创新。作品不是单一拍照 App，也不是单一云台设备，而是把手机端、设备运行时、业务后端和管理后台组合成可运行的物联网拍摄系统，实现从拍摄前模板准备、拍摄中联动控制到拍摄后 AI 分析和历史沉淀的完整闭环。",
            "手机主画面与设备转动分离。当前实现以手机本地 CameraPreview 作为用户主预览和最终媒体来源，设备侧专注于接收帧、检测、跟踪和云台控制，减少设备回传画面作为主预览带来的延迟、方向和保存链路复杂度。",
            "结构化模板构图机制。系统将模板图片转化为 bbox、head_box、points、segments、shoulder_center、face_center 等可复用结构，便于移动端显示、设备端对齐和后端保存，而不是仅依赖静态参考图。",
            "AI Provider 统一治理。后端集中管理 OpenAI-compatible 视觉模型 Provider，并通过套餐 feature_flags 控制可用模型，使 AI 能力可配置、可替换、可记录，避免把生产 Key 分散到手机或树莓派端。",
            "硬件驱动抽象。设备端同时支持 mock 与 ttl_bus 两类驱动，开发时可无硬件调试，实物演示时可切换到 TTL 总线舵机，提升工程可维护性和课堂/比赛演示稳定性。",
        ],
    )
    add_table(
        doc,
        ["创新点", "实现依据", "实际价值"],
        [
            ["四端协同", "移动端、设备端、后端、后台、数据库职责分离", "兼顾实时控制、业务管理和后续运营扩展。"],
            ["手机主预览", "移动端 README 与设备联动页面均强调手机负责画面", "保证用户所见与最终保存媒体一致。"],
            ["模板数据契约", "template_data、构图评分引擎、移动端 overlay 结构", "便于复拍、构图评分和跨端坐标统一。"],
            ["AI 后端治理", "Provider 配置、套餐 feature_flags、AI 任务记录", "便于模型切换、权限分级、成本控制和结果追踪。"],
            ["云台驱动抽象", "驱动配置支持 mock 与 ttl_bus", "降低无硬件开发成本，支持真实舵机联调。"],
        ],
    )
    add_caption(doc, "表 3  创新点与实现依据")


def add_practical(doc: Document) -> None:
    add_heading(doc, "五、实用点", 1)
    add_para(
        doc,
        "作品的实用性主要体现在“普通用户可用、开发演示可跑、后续扩展有接口”三个层面。对普通用户而言，系统能通过模板构图、手势倒计时和自动跟随减少反复试拍；对开发和教学场景而言，mock 驱动、本机运行时和清晰接口使无硬件环境也能演示主要流程；对后续产品化而言，后端 AI Provider、套餐 feature_flags 和管理后台为运营扩展留下空间。",
    )
    add_table(
        doc,
        ["应用场景", "可用能力", "实际价值"],
        [
            ["个人自拍/合影", "手机主预览、模板构图、手势倒计时、照片保存", "减少对他人帮拍的依赖，提高构图稳定性。"],
            ["旅行记录/校园活动", "设备联动、自动跟随、历史记录、AI 分析", "适合固定机位或轻量云台辅助拍摄。"],
            ["短视频/直播辅助", "云台控制、目标跟随、背景锁定、AI 多角度候选分析", "帮助创作者快速找到更合适的画面角度。"],
            ["教学与竞赛演示", "mock 与 ttl_bus 双驱动、REST/WebSocket 接口、管理后台", "便于课堂展示物联网软硬协同与多端工程结构。"],
            ["后续商业化验证", "套餐、Provider 配置、AI 任务记录、后台运营", "可扩展为分级服务、模板市场或自助拍摄终端。"],
        ],
    )
    add_caption(doc, "表 4  应用场景与价值")
    add_para(
        doc,
        "同时，本说明书对当前尚未完全固化的内容采取保守表述。例如，文档不虚构最终量产外壳、长期稳定性指标或未实测的硬件性能数据；对于真实树莓派、最终舵机负载、供电安全和现场网络抗干扰能力，仍建议在后续实物联调中补充专项测试记录。",
    )


def add_summary(doc: Document, anonymous: bool) -> None:
    add_heading(doc, "六、总结", 1)
    add_para(
        doc,
        "“云影随行”围绕手机摄影中“看不准、站不好、调不稳、选不出”的实际痛点，构建了一个由移动端、设备运行时、业务后端和管理后台组成的物联网智能拍摄系统。当前项目已经形成清晰工程边界：手机负责主画面与最终媒体，设备负责本地检测与云台转动，服务器负责 AI 与业务数据，后台负责管理配置。作品的核心价值在于把相机交互、视觉检测、云台执行、AI 分析和后台治理连接成可演示、可维护、可迭代的完整链路；后续可继续完善真实硬件外壳、供电安全、长时间稳定性测试、多设备管理和 AI 创作能力，向文旅自助拍摄、校园活动记录和轻量化内容创作辅助设备等方向延伸。",
    )


def scrub_core_properties(doc: Document, anonymous: bool) -> None:
    props = doc.core_properties
    props.author = "" if anonymous else STUDENTS
    props.last_modified_by = "" if anonymous else STUDENTS
    props.title = "参赛作品说明书"
    props.subject = PROJECT_NAME
    props.keywords = "" if anonymous else "云影随行; 参赛作品说明书"
    props.comments = ""
    props.category = ""


def build_doc(path: Path, anonymous: bool, arch_img: Path, flow_img: Path) -> None:
    doc = Document()
    style_document(doc)
    add_cover(doc, anonymous)
    add_contents(doc)
    add_intro(doc)
    add_design(doc, arch_img, flow_img)
    add_team(doc, anonymous)
    add_innovation(doc)
    add_practical(doc)
    add_summary(doc, anonymous)
    scrub_core_properties(doc, anonymous)
    doc.save(path)


def main() -> None:
    ensure_dir()
    arch_img = make_architecture_diagram()
    flow_img = make_flow_diagram()
    build_doc(SIGNED_DOCX, False, arch_img, flow_img)
    build_doc(ANON_DOCX, True, arch_img, flow_img)
    print(SIGNED_DOCX)
    print(ANON_DOCX)


if __name__ == "__main__":
    main()
