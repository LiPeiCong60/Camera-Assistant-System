from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables" / "competition_doc_2026"
DOCX_PATH = OUT_DIR / "智能拍照助手-物联网应用类作品技术文档-2026版.docx"
ARCH_IMG = OUT_DIR / "system_architecture.png"
FLOW_IMG = OUT_DIR / "core_flow.png"
LOGO_PATH = ROOT / "logo.png"

ACCENT = RGBColor(0x00, 0x00, 0x00)
ACCENT_LIGHT = "DCEFF4"
BORDER = "8FB9C6"
TEXT = RGBColor(0x00, 0x00, 0x00)


def ensure_dir() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)


def east_asia_font(run, name: str, size: float, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: str = "6") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def style_table(table, *, header_fill: str = ACCENT_LIGHT) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    east_asia_font(run, "宋体", 10.5)
        if r == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        east_asia_font(run, "宋体", 10.5, bold=True, color=ACCENT)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for style_name, font_name, size in (
        ("Title", "宋体", 18),
        ("Heading 1", "宋体", 14),
        ("Heading 2", "宋体", 14),
        ("Heading 3", "宋体", 14),
    ):
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT


def add_para(doc: Document, text: str, *, style: str | None = None, align=WD_ALIGN_PARAGRAPH.LEFT,
             first_line_chars: float = 2.0, after: float = 6.0, before: float = 0.0,
             font="宋体", size=10.5, bold=False, color: RGBColor | None = None) -> None:
    para = doc.add_paragraph(style=style)
    para.alignment = align
    para.paragraph_format.first_line_indent = Cm(0.74 * first_line_chars / 2.0) if first_line_chars else Cm(0)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.line_spacing = 1.45
    run = para.add_run(text)
    east_asia_font(run, font, size, bold=bold, color=color or TEXT)


def add_heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_paragraph(style=f"Heading {level}")
    para.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    para.paragraph_format.space_after = Pt(4)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    east_asia_font(run, "宋体", 14, bold=True, color=ACCENT)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.6)
        para.paragraph_format.first_line_indent = Cm(-0.4)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.35
        run1 = para.add_run("• ")
        east_asia_font(run1, "宋体", 10.5, bold=True, color=ACCENT)
        run2 = para.add_run(item)
        east_asia_font(run2, "宋体", 10.5, color=TEXT)


def draw_rounded_box(draw: ImageDraw.ImageDraw, xy, fill, outline, radius=18, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def load_font(size: int, *, bold: bool = False):
    candidates = [
        ("C:/Windows/Fonts/simsun.ttc", False),
        ("C:/Windows/Fonts/simhei.ttf", bold),
        ("C:/Windows/Fonts/msyhbd.ttc", bold),
        ("C:/Windows/Fonts/msyh.ttc", False),
    ]
    for path, requires_bold in candidates:
        if Path(path).exists():
            try:
                return ImageFont.truetype(path, size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        test = current + ch
        if draw.textlength(test, font=font) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_text_block(draw: ImageDraw.ImageDraw, box, title: str, body: str) -> None:
    x1, y1, x2, y2 = box
    title_font = load_font(28, bold=True)
    body_font = load_font(22)
    draw_rounded_box(draw, box, fill="#F7FBFC", outline="#5E9FB1")
    draw.text((x1 + 22, y1 + 16), title, font=title_font, fill="#000000")
    lines = wrap_text(draw, body, body_font, x2 - x1 - 44)
    y = y1 + 60
    for line in lines[:4]:
        draw.text((x1 + 22, y), line, font=body_font, fill="#000000")
        y += 30


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color="#357A8A", width=5):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        draw.polygon([(ex, ey), (ex - 16 * sign, ey - 9), (ex - 16 * sign, ey + 9)], fill=color)
    else:
        sign = 1 if ey > sy else -1
        draw.polygon([(ex, ey), (ex - 9, ey - 16 * sign), (ex + 9, ey - 16 * sign)], fill=color)


def generate_diagrams() -> None:
    title_font = load_font(34, bold=True)
    note_font = load_font(22)

    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 36), "智能拍照助手总体架构图", font=title_font, fill="#0E4F62")
    draw.text((60, 84), "源码呈现为“移动端 + 设备运行时 + 后端平台 + 管理后台”四端协同结构", font=note_font, fill="#4A4A4A")
    draw_text_block(draw, (70, 180, 430, 355), "移动端 Flutter", "登录、模板管理、拍摄会话、历史缓存、设备联动、局域网预览与 WebRTC。")
    draw_text_block(draw, (600, 140, 1010, 340), "设备运行时 FastAPI", "视频采集、人体检测、模板构图评估、自动跟踪、手势抓拍、背景锁位、局域网接口。")
    draw_text_block(draw, (1140, 180, 1525, 355), "后端业务平台 FastAPI", "用户、套餐、订阅、模板、拍摄记录、AI 任务、AI Provider 配置与静态资源托管。")
    draw_text_block(draw, (1140, 540, 1525, 740), "管理后台 Vue 3", "运营管理用户、设备、套餐、推荐模板、拍摄记录、AI 任务和 Provider。")
    draw_text_block(draw, (600, 540, 1010, 740), "本地硬件与视觉链路", "摄像头/手机推流、云台/舵机、串口控制、OpenCV、MediaPipe、可选 YOLO 检测。")
    draw_text_block(draw, (70, 540, 430, 740), "结构化数据与文件", "PostgreSQL 业务表、uploads 上传目录、captures 本地抓拍目录、模板库与 JSON 索引。")
    draw_arrow(draw, (430, 265), (600, 240))
    draw_arrow(draw, (1010, 240), (1140, 265))
    draw_arrow(draw, (1335, 355), (1335, 540))
    draw_arrow(draw, (805, 340), (805, 540))
    draw_arrow(draw, (430, 640), (600, 640))
    draw_arrow(draw, (1010, 640), (1140, 640))
    img.save(ARCH_IMG)

    img2 = Image.new("RGB", (1600, 920), "#FFFFFF")
    draw2 = ImageDraw.Draw(img2)
    draw2.text((60, 36), "核心业务闭环流程图", font=title_font, fill="#0E4F62")
    draw2.text((60, 84), "从模板生成到设备构图、抓拍和 AI 分析的关键路径", font=note_font, fill="#4A4A4A")
    boxes = [
        (70, 220, 300, 360, "1. 模板准备", "移动端上传原图或提交关键点，后端/端侧生成 template_data。"),
        (365, 220, 595, 360, "2. 会话创建", "移动端创建拍摄会话，可选择设备与模板，并进入联动模式。"),
        (660, 220, 890, 360, "3. 设备取流", "设备端打开单会话上下文，启动视频、检测器、云台与状态机。"),
        (955, 220, 1185, 360, "4. 实时构图", "FrameProcessor 评估模板匹配、稳定目标与 ready 状态。"),
        (1250, 220, 1480, 360, "5. 触发抓拍", "手动快门、手势倒计时或 AI 找角度任务触发本地抓拍。"),
        (365, 560, 595, 700, "6. 结果入库", "手机图像上传后端生成 captures，设备抓拍写入本地目录。"),
        (660, 560, 890, 700, "7. AI 评估", "后端按套餐选择 AI Provider，设备端可本地执行背景分析。"),
        (955, 560, 1185, 700, "8. 推荐反馈", "输出构图摘要、最佳图片、推荐 pan/tilt 增量或背景锁位框。"),
        (1250, 560, 1480, 700, "9. 历史沉淀", "拍摄记录、AI 任务、模板与缓存形成可回看、可运营、可扩展的数据闭环。"),
    ]
    for box in boxes:
        draw_text_block(draw2, box[:4], box[4], box[5])
    draw_arrow(draw2, (300, 290), (365, 290))
    draw_arrow(draw2, (595, 290), (660, 290))
    draw_arrow(draw2, (890, 290), (955, 290))
    draw_arrow(draw2, (1185, 290), (1250, 290))
    draw_arrow(draw2, (1365, 360), (1365, 500))
    draw_arrow(draw2, (1250, 630), (1185, 630))
    draw_arrow(draw2, (955, 630), (890, 630))
    draw_arrow(draw2, (660, 630), (595, 630))
    img2.save(FLOW_IMG)


def add_cover(doc: Document) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.15))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("中国大学生计算机设计大赛")
    east_asia_font(run, "宋体", 18, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("物联网应用类作品技术文档")
    east_asia_font(run, "宋体", 16, bold=True, color=TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("智能拍照助手（Camera Assistant）")
    east_asia_font(run, "宋体", 20, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("基于移动端、设备运行时与云端管理协同的智能拍摄系统")
    east_asia_font(run, "宋体", 10.5, color=TEXT)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(9.2)
    rows = [
        ("作品编号", "待填写"),
        ("作品名称", "智能拍照助手（Camera Assistant）"),
        ("版本编号", "V1.0"),
        ("填写日期", "2026年4月29日"),
    ]
    for i, (k, v) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.text = k
        c2.text = v
    style_table(table)

    add_para(
        doc,
        "文档说明：本文件依据用户提供的“2026版物联网应用类作品技术文档”章节要求重新编制，内容从项目源码实现出发重建，未沿用仓库现有说明文档表述。",
        first_line_chars=0,
        after=8,
        before=18,
        font="宋体",
        size=10.5,
    )
    doc.add_page_break()


def add_overview_section(doc: Document) -> None:
    add_heading(doc, "1. 作品概述", 1)
    add_para(doc, "智能拍照助手是一套面向人物拍摄场景的物联网应用系统，围绕“拍什么、怎么构图、由谁控制设备、拍完如何筛选与分析”四个核心问题展开设计。系统由移动端、设备运行时、后端业务平台和管理后台四部分构成，既支持纯手机拍摄，也支持手机与本地云台设备协同拍摄。")
    add_para(doc, "从源码实现看，本作品的核心能力包括：用户登录与套餐管理、姿态模板生成与推荐模板下发、设备端实时人体检测与云台跟踪、模板引导构图、手势触发抓拍、拍后 AI 构图分析、背景锁位、批量优选和运营后台治理。它不是单一的相机 App，而是一条从前端拍摄到设备控制再到云端评估与管理的数据闭环。")
    add_bullets(
        doc,
        [
            "设计思路创新性：把“后台业务平台”和“本地设备运行时”拆成双核心，既保留云端数据治理能力，又保证局域网实时控制闭环。",
            "技术创新性：以结构化 template_data 为跨端契约，让移动端、后端和设备端围绕同一套姿态模板数据协同工作。",
            "硬件创新性：设备控制层抽象了 mock 与 ttl_bus 两类舵机驱动，便于从开发仿真过渡到真实硬件。",
            "应用创新性：把模板构图、自动跟踪、手势抓拍、背景锁位和 AI 选图合并到一个统一拍摄流程中。",
        ],
    )


def add_requirements_section(doc: Document) -> None:
    add_heading(doc, "2. 需求分析", 1)
    add_para(doc, "人物自拍、旅拍和小型内容创作常见痛点包括：拍摄者无法同时兼顾取景与出镜、普通手机相机缺少稳定的实时构图反馈、纯云台设备通常缺少模板指导和后续图片评估能力，而单纯的 AI 修图应用又无法反向驱动拍摄动作。本作品的目标是在保证交互门槛较低的前提下，把“拍前指导、拍中控制、拍后分析”串成一体化体验。")
    add_para(doc, "系统面向的主要用户包括个人内容创作者、校园团队、小型影像工作室、景区或展厅的自助拍摄场景，以及具备云台或摄像头硬件基础的轻量化创业实践团队。源码中可以看到，它既支持 mobile_only 模式，也支持 device_link、AUTO_TRACK、SMART_COMPOSE 等联动模式，说明产品设计从一开始就兼顾了纯软件和软硬协同两条落地路径。")
    add_para(doc, "主要需求可以归纳为三类：一是模板化构图能力，要求用户能够快速生成、保存、选择并复用姿态模板；二是实时控制能力，要求设备端能进行取流、检测、跟踪、手势抓拍和预览回传；三是业务治理能力，要求后端提供用户、套餐、设备、拍摄记录和 AI 任务的统一管理。")

    table = doc.add_table(rows=4, cols=4)
    data = [
        ("比较维度", "普通手机相机/自拍杆", "纯云台跟拍设备", "本作品"),
        ("拍前构图指导", "主要依赖人工经验", "通常只有取景和跟随", "支持模板生成、推荐模板和实时构图评分"),
        ("拍中自动能力", "手动为主", "有自动跟随但少有姿态约束", "支持自动跟踪、手势倒计时、背景锁位和 AI 找角度"),
        ("拍后数据治理", "照片分散在本地", "多停留在设备侧", "后端记录拍摄会话、图片、AI 任务和套餐能力"),
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = value
    style_table(table)


def add_tech_solution_section(doc: Document) -> None:
    add_heading(doc, "3. 技术方案", 1)
    add_heading(doc, "3.1 系统总体设计", 2)
    add_para(doc, "系统采用四端协同架构。移动端负责用户交互、拍摄入口和设备联动；设备运行时负责低时延视觉闭环与本地控制；后端业务平台负责账户、套餐、模板、拍摄记录与 AI 任务治理；管理后台负责运营配置。该方案将“实时性要求高的能力”尽量放在本地，将“需要长期存储和运营的能力”放在服务端。", first_line_chars=2)
    doc.add_picture(str(ARCH_IMG), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "3.2 硬件组成与来源", 2)
    add_bullets(
        doc,
        [
            "移动终端：Flutter 客户端运行在手机上，负责摄像头拍摄、模板管理和设备联动控制。",
            "设备侧主机：从源码配置项与依赖看，设备运行时可部署在具备 Python、OpenCV 和 MediaPipe 能力的本地主机或树莓派类设备。",
            "视觉输入：支持本地摄像头索引、RTSP/HTTP 视频流，也支持 mobile_push 模式由手机直接推送 NV21/JPEG 帧。",
            "云台控制：控制层支持 mock 与 ttl_bus 两种舵机驱动抽象，串口端口、角度限制和速度行为均由配置控制。",
            "服务端与数据库：后端基于 FastAPI 和 PostgreSQL，负责账号、模板、设备、拍摄会话和 AI 任务管理。",
        ],
    )

    add_heading(doc, "3.3 软件设计与关键模块", 2)
    table = doc.add_table(rows=5, cols=3)
    data = [
        ("模块", "技术栈", "职责"),
        ("backend", "FastAPI + SQLAlchemy + PostgreSQL", "账户认证、套餐订阅、模板存储、拍摄记录、AI 任务与 Provider 配置"),
        ("device_runtime", "FastAPI + OpenCV + MediaPipe + WebSocket/WebRTC", "视频取流、实时检测、云台控制、模板构图、手势抓拍、本地 AI 编排"),
        ("mobile_client", "Flutter", "用户登录、拍照上传、模板创建、历史缓存、设备控制与预览"),
        ("admin_web", "Vue 3 + Pinia + Element Plus", "用户、设备、套餐、推荐模板、AI Provider 和统计管理"),
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = value
    style_table(table)

    add_heading(doc, "3.4 接口通用性与可扩展性", 2)
    add_bullets(
        doc,
        [
            "后端接口采用 REST 风格，并用统一 success/message/data 结构表达常规成功响应。",
            "设备端接口在 REST 之外补充了 WebSocket 预览、手机推流协议和 WebRTC offer/answer 协商，兼顾控制与实时媒体。",
            "template_data 作为结构化模板契约，允许模板来源于后端识别、移动端本地识别或设备端导入。",
            "AI 能力通过 Provider 配置与本地接口抽象扩展，后端面向多 Provider 配置，设备端面向 AIPhotoAssistant 实现扩展。",
        ],
    )

    add_heading(doc, "3.5 代码规范与安全性", 2)
    add_para(doc, "源码分层较清晰：后端按路由、依赖、服务、仓储、模型、模式对象划分；设备运行时按 API、会话管理、视觉、控制、模板与 AI 编排划分。后端认证使用自定义 HMAC 签名 Bearer Token，设备端下载抓拍文件时限制路径必须位于 captures 目录下，云台运动命令通过模式和偏移量进行约束，这些实现体现出一定的边界意识与安全设计。")

    add_heading(doc, "3.6 自主知识产权技术说明", 2)
    add_para(doc, "从代码实现层面看，本作品的核心业务逻辑为自主开发，包括：模板数据结构设计与多端复用机制、设备端单会话状态机、实时帧处理与构图评分逻辑、手势倒计时抓拍链路、AI 找角度与背景锁位任务编排、套餐到 AI Provider 的路由规则等。是否已申请软件著作权、专利或其他登记，应以实际申报材料为准；本文仅说明源码中的自主实现内容。")
    add_para(doc, "企业命题作品说明：本项目源码未体现特定企业命题设备或 SDK 的强绑定，本项不适用。", first_line_chars=0)


def add_impl_section(doc: Document) -> None:
    add_heading(doc, "4. 方案实现", 1)
    add_heading(doc, "4.1 后端业务实现", 2)
    add_para(doc, "后端入口位于 backend/app/main.py，负责装配 FastAPI 应用、CORS、静态 uploads 挂载、统一异常处理和 /api 路由。认证逻辑位于 AuthService 与 app/core/auth.py：登录成功后把 user_code、role、iat、exp 编码为 JSON 负载，再使用 HMAC-SHA256 生成签名，构成轻量 Bearer Token。业务数据以 PostgreSQL 持久化，schema.sql 中定义了 users、plans、user_subscriptions、devices、templates、capture_sessions、captures、ai_tasks、ai_provider_configs 等核心表。")
    add_para(doc, "移动端相关核心服务由 MobileService 实现，支持用户模板创建、拍摄会话创建、图片上传、AI 构图分析、背景分析、批量优选和历史记录查询。值得注意的是，套餐 feature_flags 中的 default_ai_provider_code 与 available_ai_provider_codes 会直接影响 AI 路由策略，使平台具备“按套餐开放不同 AI 能力”的运营扩展空间。")

    add_heading(doc, "4.2 移动端实现", 2)
    add_para(doc, "Flutter 客户端由 CameraAssistantApp 启动，AuthController 负责会话恢复与用户状态维护。拍摄页负责相机拍照、图片上传与捕获记录创建，首页与历史页承担模板与拍摄记录展示，设备联动页则聚合设备状态轮询、模板选择、控制命令发送、WebRTC 或局域网预览和手势/AI 开关控制。客户端还实现了本地缓存机制，在后端暂时不可用时仍可回退展示部分历史与模板数据。")

    add_heading(doc, "4.3 设备运行时实现", 2)
    add_para(doc, "设备运行时的真实入口是 device_runtime/api/app.py，核心控制枢纽是 api/session_manager.py。系统采用单会话架构：打开新会话时会先关闭旧会话，再初始化视频源、检测器、云台控制器、模板服务、AI 编排器与运行时状态对象。随后 frame_loop 持续读取最新帧，把检测结果送入 FrameProcessor，综合输出稳定目标、模板构图得分、手势 ready 状态和跟踪命令。")
    add_para(doc, "在 SMART_COMPOSE 模式下，FrameProcessor 会调用 TemplateComposeEngine.evaluate() 计算 pose_score、compose_score 和 total_score，并根据 ready 状态决定是否允许手势抓拍；在 AUTO_TRACK 模式下，TrackingController 根据目标偏差生成 pan/tilt 命令；在 AI 背景锁位模式下，运行时会抑制持续自动跟踪，以保持稳定背景构图。")
    doc.add_picture(str(FLOW_IMG), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "4.4 管理后台实现", 2)
    add_para(doc, "管理后台基于 Vue 3、Pinia 与 Element Plus 实现，接口集中封装在 admin_web/src/api/admin.js。后台可管理用户、套餐、设备、推荐模板、拍摄记录、AI 任务和 AI Provider 配置，并提供概览统计。其核心价值不只是 CRUD，而是通过套餐 feature_flags 与 Provider 配置，把用户可见能力、AI 成本控制和内容运营统一到一套平台配置机制中。")

    add_heading(doc, "4.5 接口与数据实现细节", 2)
    table = doc.add_table(rows=6, cols=3)
    data = [
        ("接口层", "典型路径", "实现作用"),
        ("后端移动接口", "/api/mobile/*", "登录注册、模板、会话、图片上传、AI 分析与历史记录"),
        ("后端管理接口", "/api/admin/*", "用户、套餐、设备、推荐模板、AI Provider 与统计"),
        ("设备控制接口", "/api/device/control/*", "云台控制、模式切换、回中和跟随模式设定"),
        ("设备媒体接口", "/api/device/preview-ws、/api/device/webrtc/offer", "预览回传、手机视频推流与双向媒体协商"),
        ("模板/AI 内部契约", "template_data、AI JSON", "模板跨端复用与 AI 结果结构化落地"),
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = value
    style_table(table)

    add_heading(doc, "4.6 关键实现亮点", 2)
    add_bullets(
        doc,
        [
            "模板生成既支持后端读取图片识别姿态，也支持移动端本地识别后直接提交 template_data。",
            "设备端手势抓拍采用状态机与倒计时机制，降低瞬时识别抖动造成的误触发。",
            "AI 找角度与背景锁位都以异步任务形式运行，不会阻塞主视觉循环。",
            "后台通过套餐 feature_flags 控制可用 AI Provider，为后续商业化分级提供了基础。",
        ],
    )


def add_test_section(doc: Document) -> None:
    add_heading(doc, "5. 测试报告", 1)
    add_para(doc, "由于当前任务以源码重建文档为主，本文测试报告以“代码实现符合性 + 可执行用例结果 + 待实机指标”三部分构成，不对源码中没有量化数据支撑的指标进行虚构。能直接运行的用例写入真实结果，涉及硬件、摄像头、串口或额外媒体依赖的指标明确标注为待联调或待实测。")

    add_heading(doc, "5.1 可执行测试结果", 2)
    table = doc.add_table(rows=3, cols=4)
    data = [
        ("测试项", "执行方式", "结果", "说明"),
        ("后端拍摄类型测试", "python -m unittest backend.tests.test_capture_type", "通过", "实际运行 1 个用例，通过，耗时约 0.053s。"),
        ("设备端手势覆盖测试", "python -m unittest device_runtime.tests.test_overlay_gesture", "受环境依赖阻塞", "导入阶段缺少 aiortc 依赖，说明当前环境未安装完整媒体组件，不代表手势逻辑本身失败。"),
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = value
    style_table(table)

    add_heading(doc, "5.2 功能符合性检查", 2)
    table = doc.add_table(rows=8, cols=4)
    data = [
        ("功能模块", "设计目标", "源码符合性", "结论"),
        ("用户认证", "支持注册、登录与会话恢复", "已实现移动端登录注册、Bearer Token 与管理员登录", "符合"),
        ("模板系统", "支持模板生成、保存、推荐与选择", "后端、移动端、设备端均存在完整模板链路", "符合"),
        ("设备会话", "支持开启、关闭和状态查询", "SessionManager 负责单会话上下文与状态快照", "符合"),
        ("自动跟踪", "支持 MANUAL/AUTO_TRACK/SMART_COMPOSE", "模式枚举、控制接口与帧处理逻辑完整", "符合"),
        ("手势抓拍", "支持稳定触发和倒计时", "手势状态机、倒计时和结果状态输出已实现", "符合"),
        ("AI 分析", "支持拍照分析、背景分析、批量优选", "后端与设备端均提供对应任务入口", "符合"),
        ("实时媒体", "支持预览回传与手机推流", "WebSocket、frame upload 与 WebRTC 路径均已实现", "需联调验证"),
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = value
    style_table(table)

    add_heading(doc, "5.3 待实机量化指标", 2)
    add_bullets(
        doc,
        [
            "端到端拍摄时延：需要在真实手机、局域网与云台设备条件下进行秒级和帧级测量。",
            "人体检测帧率与构图稳定度：与摄像头分辨率、树莓派性能档位和模型配置密切相关，需要分档测试。",
            "云台角度控制精度与重复定位误差：需要基于真实串口舵机和标定板进行实测。",
            "复杂背景、弱光与多人场景的鲁棒性：需要通过实拍数据集补充回归测试。",
        ],
    )


def add_prospect_section(doc: Document) -> None:
    add_heading(doc, "6. 应用前景", 1)
    add_para(doc, "从应用落地看，智能拍照助手适用于个人内容创作、校园活动记录、轻量摄影工作室、景区自拍点、展馆互动拍摄和小型智能拍照亭等场景。它与传统手机相机相比，优势在于模板化构图指导与设备协同；与单一云台设备相比，优势在于把拍摄前、中、后的数据链条打通；与纯后期 AI 应用相比，优势在于能把 AI 建议前置到拍摄动作本身。")
    add_para(doc, "从创业实践延展看，系统已具备继续产品化的基础：后端有套餐和 AI Provider 路由，便于做分级订阅；设备端有视频源和驱动抽象，便于适配更多硬件；移动端与后台已形成最小闭环，后续可继续扩展多设备协同、对象存储、更多 AI 模型供应商、拍摄模板市场和 SaaS 管理能力。")


def add_references_section(doc: Document) -> None:
    add_heading(doc, "7. 参考文献", 1)
    refs = [
        "[1] Camera Assistant 项目组. backend/app/main.py[Z]. 2026.",
        "[2] Camera Assistant 项目组. backend/app/services/mobile_service.py[Z]. 2026.",
        "[3] Camera Assistant 项目组. backend/app/services/ai_provider_service.py[Z]. 2026.",
        "[4] Camera Assistant 项目组. device_runtime/api/session_manager.py[Z]. 2026.",
        "[5] Camera Assistant 项目组. device_runtime/services/frame_processor.py[Z]. 2026.",
        "[6] Camera Assistant 项目组. mobile_client/lib/features/device_link/device_link_page.dart[Z]. 2026.",
        "[7] Camera Assistant 项目组. admin_web/src/api/admin.js[Z]. 2026.",
        "[8] Camera Assistant 项目组. database/schema.sql[Z]. 2026.",
    ]
    for ref in refs:
        add_para(doc, ref, first_line_chars=0, after=2, font="宋体", size=10.5)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("智能拍照助手技术文档")
    east_asia_font(run, "宋体", 9, color=RGBColor(0x00, 0x00, 0x00))


def build_doc() -> None:
    ensure_dir()
    generate_diagrams()
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_overview_section(doc)
    add_requirements_section(doc)
    add_tech_solution_section(doc)
    add_impl_section(doc)
    add_test_section(doc)
    add_prospect_section(doc)
    add_references_section(doc)
    add_footer(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
